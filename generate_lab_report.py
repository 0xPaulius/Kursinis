"""
Laboratorinio darbo nr. 1 ataskaitos generatorius.
Naudoja tikrus sablonu failus.

Autorius: Paulius Turauskas
"""
from __future__ import annotations

import copy
from datetime import date
from pathlib import Path

import openpyxl
from openpyxl import load_workbook
from openpyxl.utils import get_column_letter
from openpyxl.styles import Alignment, Font, PatternFill, Border, Side

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import lxml.etree as etree

# ---------------------------------------------------------------------------
# Paths
# ---------------------------------------------------------------------------
DL = Path("/dl")
OUT = Path("/out")
TC_TEMPLATE   = DL / "test case template.xlsx"
DEF_TEMPLATE  = DL / "Defect description template.xlsx"
STUDENT       = "Paulius Turauskas"
TEST_DATE     = "2026-04-04"
BASE_NAME     = "Paulius_Turauskas_ISK_group_Lab_no_1"

# ---------------------------------------------------------------------------
# Test case data
# ---------------------------------------------------------------------------
TEST_CASES = [
    {
        "id": "TC-001",
        "name": "Prisijungimas su teisingais kredencialais",
        "technique": "Equivalence Partitioning",
        "overall": "Pass",
        "prerequisites": [
            "Logu analizės sistema paleista (docker compose up -d)",
            "Narsykle atidaryta",
            "Vartotojas nera prisijunges (sessionStorage tuscias)",
        ],
        "test_data": [
            "URL: http://localhost:8080/login",
            "Vartotojo vardas: admin",
            "Slaptazodis: admin123",
        ],
        "steps": [
            ("1", "Atidaryti http://localhost:8080/login",
             "Rodoma prisijungimo forma su laukais Vartotojo vardas ir Slaptazodis",
             "Prisijungimo forma rodoma", "Pass"),
            ("2", "Lauke [Vartotojo vardas] ivesti: admin",
             "Laukelis priima ivesti",
             "Tekstas admin iveistas", "Pass"),
            ("3", "Lauke [Slaptazodis] ivesti: admin123",
             "Laukelis priima ivesti; simboliai slepiami (taskai)",
             "Slaptazodis iveistas, simboliai slepiami", "Pass"),
            ("4", "Paspausti mygtuka [Prisijungti]",
             "Sistema autentifikuoja vartotoja ir nukreipia i /",
             "Nukreipta i http://localhost:8080/, rodomas Apzvalgos puslapis", "Pass"),
            ("5", "Patikrinti narsykles sessionStorage",
             "Saugomi raktai auth_token ir auth_user",
             "sessionStorage.getItem('auth_token') grazina ne null reiksme", "Pass"),
        ],
    },
    {
        "id": "TC-002",
        "name": "Prisijungimas su neteisingu slaptazodziu",
        "technique": "Negative testing / error handling",
        "overall": "Pass",
        "prerequisites": [
            "Logu analizės sistema paleista",
            "Narsykle atidaryta",
            "Vartotojas nera prisijunges",
        ],
        "test_data": [
            "URL: http://localhost:8080/login",
            "Vartotojo vardas: admin",
            "Slaptazodis: neteisinga123 (klaidingas)",
        ],
        "steps": [
            ("1", "Atidaryti http://localhost:8080/login",
             "Rodoma prisijungimo forma",
             "Prisijungimo forma rodoma", "Pass"),
            ("2", "Ivesti vartotoja admin ir slaptazodi neteisinga123",
             "Abu laukeliai priima ivesti",
             "Duomenys ivesti", "Pass"),
            ("3", "Paspausti [Prisijungti]",
             "Sistema grazina HTTP 401 klaida",
             "Serveris grazino HTTP 401", "Pass"),
            ("4", "Patikrinti klaidos pranesima",
             "Rodomas raudonas klaidos baneris: Neteisingas vartotojo vardas arba slaptazodis",
             "Klaidos baneris rodomas su nurodytu tekstu", "Pass"),
            ("5", "Patikrinti ar slaptazodzio laukas isvalytas ir URL lieka /login",
             "Slaptazodzio laukas tuscias; URL lieka /login",
             "Slaptazodzio laukas isvalytas; URL /login", "Pass"),
        ],
    },
    {
        "id": "TC-003",
        "name": "Prisijungimas su tusciu vartotojo vardu",
        "technique": "Negative testing / error handling",
        "overall": "Pass",
        "prerequisites": [
            "Logu analizės sistema paleista",
            "Narsykle atidaryta",
            "Vartotojas nera prisijunges",
        ],
        "test_data": [
            "URL: http://localhost:8080/login",
            "Vartotojo vardas: (tuscia)",
            "Slaptazodis: admin123",
        ],
        "steps": [
            ("1", "Atidaryti http://localhost:8080/login",
             "Rodoma prisijungimo forma",
             "Prisijungimo forma rodoma", "Pass"),
            ("2", "Palikti vartotojo vardo laukeli tuscia",
             "Laukelis tuscias",
             "Laukelis tuscias", "Pass"),
            ("3", "Lauke [Slaptazodis] ivesti: admin123",
             "Slaptazodis iveistas",
             "Slaptazodis iveistas", "Pass"),
            ("4", "Paspausti [Prisijungti]",
             "Lauko klaida: Vartotojo vardas negali buti tuscias; API kreipinys neisspiestas",
             "Klaidos tekstas po lauku rodomas; API kreipinys neisspiestas", "Pass"),
            ("5", "Patikrinti ar URL lieka /login",
             "URL lieka http://localhost:8080/login",
             "URL lieka /login", "Pass"),
        ],
    },
    {
        "id": "TC-004",
        "name": "Vartotojo vardo lauko ilgio ribines reiksmes",
        "technique": "Boundary Value",
        "overall": "Pass",
        "prerequisites": [
            "Logu analizės sistema paleista",
            "Narsykle atidaryta",
            "Vartotojas nera prisijunges",
        ],
        "test_data": [
            "Riba 1 - apatine: 0 simboliu (tuscia)",
            "Riba 2 - apatine+1: 1 simbolis: a",
            "Riba 3 - virsutine: 64 simboliai (64x a)",
            "Riba 4 - virsutine+1: 65 simboliai (65x a, draudziama)",
        ],
        "steps": [
            ("1", "Palikti vartotojo vardo lauka tuscia (0 sim.) ir paspausti [Prisijungti]",
             "Klaida: Vartotojo vardas negali buti tuscias",
             "Klaida rodoma; API neisskiestas", "Pass"),
            ("2", "Ivesti 1 simboli a vartotojo vardo laukelyje",
             "Laukelis priima 1 simboli; jokio klaidos pranesimo",
             "Simbolis a priimtas", "Pass"),
            ("3", "Ivesti 64 simbolius (virsutine riba, maxlength=64)",
             "Laukelis priima visus 64 simbolius",
             "64 simboliai ivesti; visi matomi laukelyje", "Pass"),
            ("4", "Bandyti ivesti 65-aji simboli (virsijant maxlength=64)",
             "65-asis simbolis neivedamas (HTML maxlength apribojimas)",
             "Laukelis nepriima 65-ojo simbolio; apribojimas veikia", "Pass"),
        ],
    },
    {
        "id": "TC-005",
        "name": "Prisijungimas su tarpais vietoje vartotojo vardo",
        "technique": "Negative testing / error handling",
        "overall": "Pass",
        "prerequisites": [
            "Logu analizės sistema paleista",
            "Narsykle atidaryta",
            "Vartotojas nera prisijunges",
        ],
        "test_data": [
            "Vartotojo vardas: 3 tarpai (whitespace only)",
            "Slaptazodis: admin123",
        ],
        "steps": [
            ("1", "Atidaryti /login",
             "Prisijungimo forma rodoma",
             "Forma rodoma", "Pass"),
            ("2", "Vartotojo vardo laukelyje ivesti 3 tarpus",
             "Laukelis vizualiai atrodo tuscias",
             "Tarpai ivesti; laukelis atrodo tuscias", "Pass"),
            ("3", "Ivesti teisinga slaptazodi admin123",
             "Slaptazodis iveistas",
             "Slaptazodis iveistas", "Pass"),
            ("4", "Paspausti [Prisijungti]",
             "Klaida: Vartotojo vardas negali buti tuscias (trim() pasalina tarpus)",
             "Klaidos zinute rodoma; API kreipinys neisspiestas", "Pass"),
        ],
    },
    {
        "id": "TC-006",
        "name": "Prieiga prie sistemos be prisijungimo",
        "technique": "Negative testing / error handling",
        "overall": "Pass",
        "prerequisites": [
            "Logu analizės sistema paleista",
            "Narsykle atidaryta incognito rezime (sessionStorage tuscias)",
        ],
        "test_data": [
            "URL: http://localhost:8080 (tiesioginis adresas)",
        ],
        "steps": [
            ("1", "Atidaryti nauja incognito langa narsykles",
             "Nauja sesija be jokiu issaugotu duomenu",
             "Incognito lango narsykle atidaryta", "Pass"),
            ("2", "Adreso juostoje ivesti http://localhost:8080 ir paspausti Enter",
             "Sistema aptinka kad auth_token nera; nukreipia i /login",
             "Automatiskai nukreipta i http://localhost:8080/login", "Pass"),
            ("3", "Prisijungti su teisingais duomenimis, tada atsijungti (logout mygtukas)",
             "Nukreipiama atgal i /login; sesija istrinta",
             "Nukreipta i /login; sessionStorage isvalytas", "Pass"),
            ("4", "Spausti narsykles [Atgal] mygtuka",
             "Sistema vel patikrina auth_token ir nukreipia i /login",
             "Vel nukreipta i /login", "Pass"),
        ],
    },
    {
        "id": "TC-007",
        "name": "Paieska - ekvivalentines sekcijos",
        "technique": "Equivalence Partitioning",
        "overall": "Pass",
        "prerequisites": [
            "Sistema paleista su log duomenimis",
            "Vartotojas prisijunges",
            "Anomaliju puslapyje yra bent 1 ivykis",
        ],
        "test_data": [
            "Sekcija 1 - tikslus IP: 192.168.1.100",
            "Sekcija 2 - dalinis IP: 192.168",
            "Sekcija 3 - neegzistuojantis: xyznotfound999",
        ],
        "steps": [
            ("1", "Atidaryti Anomaliju puslapi; paieska ivesti IP 192.168.1.100",
             "Chronologijoje rodomi tik to IP ivykiai; skaitliukas >0",
             "Filtruoti ivykiai rodomi; skaitliukas > 0", "Pass"),
            ("2", "Isvalyti paieska; ivesti dalini IP 192.168",
             "Rodomi visi ivykiai su IP prasidedanciu 192.168",
             "Keli ivykiai rodomi", "Pass"),
            ("3", "Isvalyti; ivesti xyznotfound999",
             "Chronologijoje rodoma: Ivykiu nera; skaitliukas 0 rezultatu",
             "Tuscias sarasas ir 0 rezultatu skaitliukas rodomas", "Pass"),
        ],
    },
    {
        "id": "TC-008",
        "name": "Paieiskos laukelio ribines reiksmes",
        "technique": "Boundary Value",
        "overall": "Pass",
        "prerequisites": [
            "Sistema paleista su log duomenimis",
            "Vartotojas prisijunges",
        ],
        "test_data": [
            "Riba 1: 0 simboliu - tuscia paieska",
            "Riba 2: 1 simbolis: a",
            "Riba 3: 100 simboliu eilute",
        ],
        "steps": [
            ("1", "Palikti paieiskos lauka tuscia (0 sim.)",
             "Rodomi visi ivykiai; rezultatu skaitliukas nematomas",
             "Visi ivykiai rodomi; skaitliukas pasleptas", "Pass"),
            ("2", "Ivesti 1 simboli: a",
             "Duomenys filtruojami; rodomas rezultatu skaitliukas",
             "Filtravimas aktyvuojamas; skaitliukas rodomas", "Pass"),
            ("3", "Ivesti 100 simboliu eilute (100x a)",
             "Paieska veikia be klaidu; rodoma 0 rezultatu",
             "Paieska atliekama; sistema nestoja; skaitliukas rodomas", "Pass"),
            ("4", "Isvalyti paieska spaudziant X mygtuka",
             "Paieska isvaloma; visi ivykiai vel rodomi; skaitliukas paslepamas",
             "Laukelis isvalytas; visi duomenys atsinaujina", "Pass"),
        ],
    },
    {
        "id": "TC-009",
        "name": "Paieska su specialiais simboliais",
        "technique": "Negative testing / error handling",
        "overall": "Pass",
        "prerequisites": [
            "Sistema paleista",
            "Vartotojas prisijunges",
        ],
        "test_data": [
            "Ivesti 1: <script>alert(xss)</script>",
            "Ivesti 2: '; DROP TABLE logs;--",
            "Ivesti 3: !@#$%^&*()",
        ],
        "steps": [
            ("1", "Paieiskos laukelyje ivesti: <script>alert(xss)</script>",
             "Tekstas rodomas kaip eilute; joks JavaScript nepaleiddziamas; 0 rezultatu",
             "Tekstas rodomas be vykdymo; 0 rezultatu; jokio alert dialogo", "Pass"),
            ("2", "Isvalyti; ivesti: '; DROP TABLE logs;--",
             "Paieska atliekama kaip paprastas tekstas; sistema veikia normaliai",
             "0 rezultatu; sistema nestojo; jokios klaidos", "Pass"),
            ("3", "Isvalyti; ivesti: !@#$%^&*()",
             "Paieska atliekama; 0 rezultatu arba atitikmuo duomenyse",
             "Sistema veikia normaliai; rezultatu sarasas pateikiamas", "Pass"),
        ],
    },
    {
        "id": "TC-010",
        "name": "Kalbos perjungimas LT ir EN",
        "technique": "Equivalence Partitioning",
        "overall": "Pass",
        "prerequisites": [
            "Sistema paleista",
            "Vartotojas prisijunges",
            "Kalba nustatyta LT (numatytoji)",
        ],
        "test_data": [
            "Sekcija 1: LT (lietuviu) - numatytoji",
            "Sekcija 2: EN (anglu) - perjungta",
        ],
        "steps": [
            ("1", "Patikrinti, kad sistema yra LT kalboje",
             "Sonines juostos meniu: Apzvalga, Anomalijos, Irenginiai, Pranesrimai",
             "Visi meniu punktai lietuviu kalba", "Pass"),
            ("2", "Paspausti EN kalbos perjungimo mygtuka",
             "Visi tekstai perjungiami i anglu kalba: Dashboard, Anomalies, Devices",
             "Tekstai perjungti i anglu kalba", "Pass"),
            ("3", "Perkrauti puslapi (F5)",
             "Kalba islieka anglu (saugoma localStorage)",
             "Po perkrovimo kalba lieka EN", "Pass"),
            ("4", "Paspausti LT ir patikrinti sugrizima",
             "Visi tekstai grizta i lietuviu kalba",
             "Tekstai grito i LT", "Pass"),
        ],
    },
    {
        "id": "TC-011",
        "name": "Atsijungimas ir sesijos pabaiga",
        "technique": "Negative testing / error handling",
        "overall": "Pass",
        "prerequisites": [
            "Sistema paleista",
            "Vartotojas prisijunges (auth_token sessionStorage)",
        ],
        "test_data": [
            "Naudojamas jau prisijunges vartotojas",
        ],
        "steps": [
            ("1", "Paspausti atsijungimo (logout) piktograma virsuje desineoje",
             "Nukreipiamas i /login; sessionStorage isvalomas",
             "Nukreipta i /login; auth_token pasalintas is sessionStorage", "Pass"),
            ("2", "Prisijungimo puslapyje bandyti ivesti http://localhost:8080",
             "Sistema patikrina auth_token (jo nera) ir nukreipia i /login",
             "Vel nukreipta i /login", "Pass"),
            ("3", "Atidaryti /login kai jau esama prisijungus (naujame lange)",
             "Sistema nukreipia i / (auth_token egzistuoja)",
             "Nukreipta i / be prisijungimo formos", "Pass"),
        ],
    },
    {
        "id": "TC-012",
        "name": "Irenginiu saraso paieska - skaitliuku tikrinimas",
        "technique": "Equivalence Partitioning",
        "overall": "Fail",
        "prerequisites": [
            "Sistema paleista su log duomenimis",
            "Vartotojas prisijunges",
            "Sistemoje yra bent 3 skirtingu IP irenginiai su ispejimais",
        ],
        "test_data": [
            "Paieska: konkretus IP, atitinkantis tik 1 irengini",
        ],
        "steps": [
            ("1", "Atidaryti Irenginiu puslapi; uzfiksuoti skaitliukus: Is viso, Kritiniai, Ispejimai",
             "Rodomi visi irenginiai ir teisingi bendri skaitliukai",
             "Skaitliukai rodo visu irenginiu kieki", "Pass"),
            ("2", "Paieiskos laukelyje ivesti konkretui IP (pvz. pirmojo irenginio IP)",
             "Lenteleje rodomas 1 atitinkantis irenginys",
             "Lenteleje rodomas 1 irenginys", "Pass"),
            ("3", "Patikrinti ar skaitliukai Is viso, Kritiniai, Ispejimai atsinaujino",
             "Skaitliukai turi rodyti filtruotus duomenis (pvz. Is viso: 1)",
             "Skaitliukai rodo nesufiltuota bendra kieki (pvz. Is viso: 10) - neatitikimas su lentele",
             "Fail"),
        ],
    },
]

DEFECT = {
    "id": "DR-001",
    "name": "Irenginiu puslapio suvestines skaitliukai neatspindi paieiskos filtro",
    "severity": "Medium",
    "priority": "Medium",
    "environment": (
        "Logu analizės sistema v1.0 | http://localhost:8080 | "
        "Narsykle: Chrome 124 / Firefox 125 | OS: Linux"
    ),
    "actions": (
        "1. Prisijungti prie sistemos (admin / admin123)\n"
        "2. Atidaryti Irenginiu puslapi (sonine juosta -> Irenginiai)\n"
        "3. Uzfiksuoti suvestines skaitliukus virsuje (Is viso, Kritiniai, Ispejimai)\n"
        "4. Paieiskos laukelyje ivesti konkretui IP adresa, atitinkanti 1 irengini\n"
        "5. Stebeti: lenteleje rodomas 1 filtruotas irenginys\n"
        "6. Stebeti: skaitliukai virsuje NESIKEICIA - rodo nesufiltuota kieki"
    ),
    "actual_result": (
        "Po paieiskos lenteleje rodomas 1 atitinkantis irenginys, taciau "
        "suvestines skaitliukai Is viso: N, Kritiniai: X, Ispejimai: Y "
        "rodo visu (nesufiltuotu) irenginiu kieki."
    ),
    "expected_result": (
        "Suvestines skaitliukai turi atspindeti tik filtruotu irenginiu kieki. "
        "Jei paieska grazina 1 irengini, skaitliukas Is viso turi rodyti 1."
    ),
    "comment": (
        "Priezastis: app/static/js/app.js, funkcija renderDevicesPage() (eilutes 818-820). "
        "Kintamasis devs = filteredDevices() naudojamas lenteles atvaizdavimui, "
        "taciau skaitliukai naudoja state.devices.length (nesufiltuotas masyvas). "
        "Sprendimas: pakeisti state.devices i devs trijose skaitliuku eilutese."
    ),
}

# ---------------------------------------------------------------------------
# XLSX: fill test case into a copied template sheet
# ---------------------------------------------------------------------------

def _set(ws, coord: str, value: str) -> None:
    ws[coord] = value
    ws[coord].alignment = Alignment(wrap_text=True, vertical="top")


def fill_tc_sheet(ws, tc: dict) -> None:
    # Header
    _set(ws, "C1", tc["id"])
    _set(ws, "F1", tc["name"])
    _set(ws, "C2", STUDENT)
    _set(ws, "C4", f"Technika: {tc['technique']}")
    _set(ws, "C6", STUDENT)
    _set(ws, "F6", TEST_DATE)
    _set(ws, "J6", tc["overall"])

    # Prerequisites (B9:D9 .. B12:D12)
    pre_rows = ["B9", "B10", "B11", "B12"]
    for i, pre in enumerate(tc["prerequisites"][:4]):
        _set(ws, pre_rows[i], pre)

    # Test data (G9:K9 .. G12:K12)
    td_rows = ["G9", "G10", "G11", "G12"]
    for i, td in enumerate(tc["test_data"][:4]):
        _set(ws, td_rows[i], td)

    # Steps (start row 18, each row: A=num, B=details, D=expected, F=actual, I=status)
    for i, step in enumerate(tc["steps"]):
        r = 18 + i
        ws[f"A{r}"] = step[0]
        _set(ws, f"B{r}", step[1])
        _set(ws, f"D{r}", step[2])
        _set(ws, f"F{r}", step[3])
        _set(ws, f"I{r}", step[4])
        ws.row_dimensions[r].height = max(20, len(step[2]) // 3 + 10)


def build_xlsx(out_path: Path) -> None:
    tc_wb  = load_workbook(TC_TEMPLATE)
    def_wb = load_workbook(DEF_TEMPLATE)
    template_ws  = tc_wb["Template"]
    defect_ws    = def_wb["Sheet1"]

    # Create a new workbook and copy each TC sheet
    out_wb = openpyxl.Workbook()
    out_wb.remove(out_wb.active)

    for tc in TEST_CASES:
        new_ws = tc_wb.copy_worksheet(template_ws)
        new_ws.title = tc["id"]
        fill_tc_sheet(new_ws, tc)

    # Remove the original template sheet
    tc_wb.remove(template_ws)

    # Rename sheets (copy_worksheet prefixes with "Template ")
    for ws in tc_wb.worksheets:
        pass  # titles already set above

    # Add defect sheet: copy template and fill
    new_def = tc_wb.copy_worksheet(defect_ws) if "Sheet1" in tc_wb.sheetnames else None
    # Easier: copy from def_wb into tc_wb
    # Since copy_worksheet only works within same workbook, manually recreate defect layout
    dws = tc_wb.create_sheet("DR-001")
    dws.column_dimensions["A"].width = 42
    dws.column_dimensions["B"].width = 62

    THIN = Side(style="thin")
    BRD  = Border(left=THIN, right=THIN, top=THIN, bottom=THIN)
    LBLUE = PatternFill("solid", fgColor="D9E1F2")

    rows = [
        (2,  "Id number",      DEFECT["id"]),
        (3,  "Name/description", DEFECT["name"]),
        (4,  "Severity",       DEFECT["severity"]),
        (5,  "Priority",       DEFECT["priority"]),
        (9,  "Environment",    DEFECT["environment"]),
        (11, "Scenario/Actions (kaip atkartoti)", DEFECT["actions"]),
        (13, "Actual Result",  DEFECT["actual_result"]),
        (15, "Expected result", DEFECT["expected_result"]),
        (17, "Comment",        DEFECT["comment"]),
    ]
    for row_num, label, value in rows:
        a = dws.cell(row=row_num, column=1, value=label)
        a.font = Font(bold=True)
        a.fill = LBLUE
        a.border = BRD
        a.alignment = Alignment(wrap_text=True, vertical="top")

        b = dws.cell(row=row_num, column=2, value=value)
        b.border = BRD
        b.alignment = Alignment(wrap_text=True, vertical="top")

        lines = value.count("\n") + 1
        dws.row_dimensions[row_num].height = max(20, 15 * lines)

    tc_wb.save(out_path)
    print(f"  XLSX: {out_path}")


# ---------------------------------------------------------------------------
# DOCX via python-docx
# ---------------------------------------------------------------------------

def _add_tc_heading(doc: Document, tc: dict) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    run = p.add_run(f"{tc['id']}: {tc['name']}")
    run.bold = True
    run.font.size = Pt(13)

    p2 = doc.add_paragraph()
    run2 = p2.add_run(f"Technika: {tc['technique']}")
    run2.italic = True
    run2.font.size = Pt(10)


def _set_cell_bg(cell, hex_color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def _bold_cell(cell) -> None:
    for para in cell.paragraphs:
        for run in para.runs:
            run.bold = True


def _cell_text(cell, text: str, bold: bool = False, size: int = 9) -> None:
    cell.text = text
    cell.paragraphs[0].paragraph_format.space_before = Pt(1)
    cell.paragraphs[0].paragraph_format.space_after = Pt(1)
    for run in cell.paragraphs[0].runs:
        run.bold = bold
        run.font.size = Pt(size)


def add_tc_tables(doc: Document, tc: dict) -> None:
    # --- Prerequisites + Test Data side-by-side (2 cols: prereqs left, testdata right) ---
    tbl = doc.add_table(rows=1 + max(len(tc["prerequisites"]), len(tc["test_data"])),
                        cols=4)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Header row
    h0 = tbl.cell(0, 0).merge(tbl.cell(0, 1))
    _cell_text(h0, "S #  Prerequisites:", bold=True)
    _set_cell_bg(h0, "BDD7EE")

    h2 = tbl.cell(0, 2).merge(tbl.cell(0, 3))
    _cell_text(h2, "S #  Test Data Requirement:", bold=True)
    _set_cell_bg(h2, "BDD7EE")

    max_rows = max(len(tc["prerequisites"]), len(tc["test_data"]))
    for i in range(max_rows):
        r = i + 1
        pre = tc["prerequisites"][i] if i < len(tc["prerequisites"]) else ""
        td  = tc["test_data"][i]     if i < len(tc["test_data"])     else ""
        _cell_text(tbl.cell(r, 0), str(i + 1))
        _cell_text(tbl.cell(r, 1), pre)
        _cell_text(tbl.cell(r, 2), str(i + 1))
        _cell_text(tbl.cell(r, 3), td)

    doc.add_paragraph()

    # --- Steps table ---
    tbl2 = doc.add_table(rows=1 + len(tc["steps"]), cols=5)
    tbl2.style = "Table Grid"
    tbl2.alignment = WD_TABLE_ALIGNMENT.LEFT

    hdrs = ["Step #", "Step Details", "Expected Results",
            "Actual Results", "Pass / Fail / Not executed / Suspended"]
    widths_cm = [1.0, 3.5, 4.0, 4.0, 2.8]
    for col_idx, (h, w) in enumerate(zip(hdrs, widths_cm)):
        cell = tbl2.cell(0, col_idx)
        _cell_text(cell, h, bold=True)
        _set_cell_bg(cell, "BDD7EE")
        cell.width = Cm(w)

    GREEN = "C6EFCE"
    RED   = "FFC7CE"
    for row_idx, step in enumerate(tc["steps"], 1):
        _cell_text(tbl2.cell(row_idx, 0), step[0])
        _cell_text(tbl2.cell(row_idx, 1), step[1])
        _cell_text(tbl2.cell(row_idx, 2), step[2])
        _cell_text(tbl2.cell(row_idx, 3), step[3])
        status_cell = tbl2.cell(row_idx, 4)
        _cell_text(status_cell, step[4], bold=True)
        _set_cell_bg(status_cell, GREEN if step[4] == "Pass" else RED)

    doc.add_page_break()


def add_defect_table(doc: Document, defect: dict) -> None:
    p = doc.add_paragraph()
    run = p.add_run("Defekto aprasymas")
    run.bold = True
    run.font.size = Pt(13)

    rows_data = [
        ("Id number",       defect["id"]),
        ("Name/description", defect["name"]),
        ("Severity",        defect["severity"]),
        ("Priority",        defect["priority"]),
        ("Environment",     defect["environment"]),
        ("Scenario/Actions (kaip atkartoti)", defect["actions"]),
        ("Actual Result",   defect["actual_result"]),
        ("Expected result", defect["expected_result"]),
        ("Comment",         defect["comment"]),
    ]
    tbl = doc.add_table(rows=len(rows_data), cols=2)
    tbl.style = "Table Grid"

    for i, (label, value) in enumerate(rows_data):
        _cell_text(tbl.cell(i, 0), label, bold=True)
        _set_cell_bg(tbl.cell(i, 0), "F4CCCC")
        _cell_text(tbl.cell(i, 1), value)
        tbl.cell(i, 0).width = Cm(5.5)
        tbl.cell(i, 1).width = Cm(10.5)


def build_docx(out_path: Path) -> None:
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin   = Cm(2)
        section.right_margin  = Cm(2)

    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(
        "Laboratorinis darbas nr. 1\n"
        "Testavimo atveju ir defektu rasymasz"
    )
    run.bold = True
    run.font.size = Pt(16)

    doc.add_paragraph()
    info_lines = [
        f"Studentas: {STUDENT}",
        "Testuojama aplikacija: Logu analizės ir anomaliju aptikimo sistema",
        "URL: http://localhost:8080",
        f"Data: {TEST_DATE}",
        "Is viso testavimo atveju: 12  |  Defektu aprasymai: 1",
    ]
    for line in info_lines:
        p = doc.add_paragraph(line)
        p.runs[0].font.size = Pt(11)

    doc.add_page_break()

    for tc in TEST_CASES:
        _add_tc_heading(doc, tc)
        add_tc_tables(doc, tc)

    add_defect_table(doc, DEFECT)

    doc.save(out_path)
    print(f"  DOCX: {out_path}")


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def main() -> None:
    print("Generating lab report...")
    build_xlsx(OUT / f"{BASE_NAME}.xlsx")
    build_docx(OUT / f"{BASE_NAME}.docx")
    print("Done.")


if __name__ == "__main__":
    main()
